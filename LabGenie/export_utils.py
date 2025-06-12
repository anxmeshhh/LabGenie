from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet
import io
import base64

def export_to_docx(experiment, file_path):
    doc = Document()
    doc.add_heading(f'Lab Record: {experiment["name"]}', 0)
    
    doc.add_heading('Aim', level=1)
    doc.add_paragraph(experiment['aim'])
    
    doc.add_heading('Theory', level=1)
    doc.add_paragraph(experiment['theory'])
    
    doc.add_heading('Procedure', level=1)
    doc.add_paragraph(experiment['procedure'])
    
    doc.add_heading('Result', level=1)
    doc.add_paragraph(experiment['result'])
    
    doc.add_heading('Graph', level=1)
    if experiment['graph']:
        graph_data = base64.b64decode(experiment['graph'])
        with open('temp_graph.png', 'wb') as f:
            f.write(graph_data)
        doc.add_picture('temp_graph.png')
    
    doc.save(file_path)

def export_to_pdf(experiment, file_path):
    doc = SimpleDocTemplate(file_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    story.append(Paragraph(f'Lab Record: {experiment["name"]}', styles['Title']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Aim', styles['Heading1']))
    story.append(Paragraph(experiment['aim'], styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Theory', styles['Heading1']))
    story.append(Paragraph(experiment['theory'], styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Procedure', styles['Heading1']))
    story.append(Paragraph(experiment['procedure'], styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph('Result', styles['Heading1']))
    story.append(Paragraph(experiment['result'], styles['Normal']))
    story.append(Spacer(1, 12))
    
    if experiment['graph']:
        graph_data = base64.b64decode(experiment['graph'])
        with open('temp_graph.png', 'wb') as f:
            f.write(graph_data)
        story.append(Image('temp_graph.png', width=400, height=300))
    
    doc.build(story)